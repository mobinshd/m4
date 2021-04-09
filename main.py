# -*- coding: utf-8 -*-
"""Copy of new book-classification.ipynb

Automatically generated by Colaboratory.

Original file is located at
    https://colab.research.google.com/drive/1SxmSSxfO0Ps4ibCNKzSA09R3_bMH4V9u
"""

#from google.colab import drive
#drive.mount('/content/drive')

"""# Library"""

#!pip install python-docx
#!pip install parsivar
#!wget https://www.dropbox.com/s/tlyvnzv1ha9y1kl/spell.zip
#!pip install langdetect

#!python --version

#!unzip spell.zip

#!mkdir /usr/local/lib/python3.7/dist-packages/parsivar/resource/spell
#!cp onegram.pckl /usr/local/lib/python3.7/dist-packages/parsivar/resource/spell/
#!cp mybigram_lm.pckl /usr/local/lib/python3.7/dist-packages/parsivar/resource/spell/

import sqlalchemy  
from sqlalchemy.ext.automap import automap_base
from sqlalchemy.orm import Session
from sqlalchemy import create_engine
from docx import Document
import os
import pandas as pd
from parsivar import Normalizer
from parsivar import Tokenizer
from parsivar import FindStems
from parsivar import SpellCheck
import matplotlib.pyplot as plt
import seaborn as sns
import pickle
import re
import numpy as np
from langdetect import detect
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
from sklearn.model_selection import train_test_split
from sklearn.feature_selection import chi2
import sklearn.metrics as metrics
import sys
from sklearn.naive_bayes import MultinomialNB
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier
import xgboost
from sys import getsizeof
import gc

"""# 1-Initial Dataset

"""
def ShowContent_doc(content):
    print(content)
    print('_________________________________________________________')

####################################################

def LanguageDetection_doc(doc):
  rng=3000 #at least lenght of str

  lan=detect(doc.content[0:rng])
  return lan


if os.path.isfile('data.pkl'):
    df = pd.read_pickle('data.pkl')
else:
    df = pd.DataFrame(data={'filename': [], 'content': [],'label': [],'labelfarsi':[]})
    for dirName, subDirList, fileList in os.walk('light data/',topdown=True):
        print(dirName)
        print(subDirList)
        print(fileList)
        exluded_files = ['124843880', '215690387', '99664244',  '107706717', '71910836']
        for file in fileList:
              filename,extension = os.path.splitext(file)             
              x= dirName.split('-',1)
              if filename in exluded_files:
                  continue
              d=Document(dirName+'/'+filename+'.docx')
              text = ''
              for p in d.paragraphs:
                text += p.text
              df=df.append(pd.DataFrame({'filename':[filename],
                                  'content':[text] ,
                                  'label': [x[0].split('/')[-1]] ,
                                  'labelfarsi':[ x[1]],
                                  'paragrapghs_length':[len(d.paragraphs)],
                                  'content_length':[len(text)]}),ignore_index=True)
        gc.collect()

    """# 2-Data Cleaning

    **Data cleaning function**
    """


    """**removing empty docs**"""

    #بعضی اسناد خالی هستند و باید حذف شوند
    t=df.loc[(df.paragrapghs_length <50) | (df.content_length<3000)]
    df.drop(t.index,inplace=True)
    df.reset_index(inplace=True)

    """**language detection**"""

    language=[]
    for i in range(len(df)):
      language.append( LanguageDetection_doc(df.loc[i]))

    df.insert(5,"language",language)

    """**grouping languages**"""

    df.groupby(['language']).size().reset_index(name='counts')

    """**none Farsi Detection**"""

    nonfarsi=df.loc[ (df['language']!='fa') ]
    nonfarsi

    """**removing non Farsi Documents**"""

    df.drop(nonfarsi.index ,inplace=True)
    df.reset_index(inplace=True)

    """# 3-Data Analysis

    **Number of articles in each category**
    """

    labellist=df.groupby(["label"]).size().reset_index(name='counts')
    labellist

    """**Plot classes**"""

    plt.figure(figsize=(12.8,6))
    sns.boxplot(data=labellist, x='label', y='counts', width=.5);

    # make a pickle from the dataframe
    df.to_pickle('data.pkl')
"""# 4-Data Preprocessing

**Preprocessing Functions**
"""

document = Document('persian.docx')

def Normalizing_doc(content):
  my_normalizer = Normalizer(statistical_space_correction=True,date_normalizing_needed=True)
  normalized_doc=my_normalizer.normalize(content,new_line_elimination= True)   
  return normalized_doc

####################################################

def SpellChecking_doc(normalized_content):
  myspell_checker = SpellCheck()
  spellchecked_doc = myspell_checker.spell_corrector(normalized_content)
  return spellchecked_doc

####################################################


def Tokenizing_doc(spellchecked_content):
    my_tokenizer = Tokenizer()
    tokens = my_tokenizer.tokenize_words(spellchecked_content)
    return tokens

####################################################

def Stemming_doc(tokens):
  stemmed_words=[]
  my_stemmer = FindStems()
  for p in tokens :
    stemmed_words.append( my_stemmer.convert_to_stem(p))
  return stemmed_words

  ####################################################
  

stopwords=[]
for i in range(len(document.paragraphs)):
   stopwords.append( document.paragraphs[i].text)

reg = re.compile(r'[a-zA-Z]|\d+(\.\d+)?|\.|\!|\؟|\'|\»|\،|\(|\)')

def remove_uselesstokens(token):
 return [s for s in token if  not (reg.match(s))]

def remove_stopwords(token):
  return [s for s in token if (s.split("&")[0]  not in stopwords)]

  #multi sense words

tk=['داشت&دار','داد&ده','آموخت&آموز','شد&شو','کرد&کن','توانست&توان','رسید&رس','شناخت&شناس','گشت&گرد','ساخت&ساز','خواست&خواه']

tk[2].split("&")[0]

lk=[s for s in tk if (s.split("&")[0]  not in stopwords)]
lk

len(stopwords)

ShowContent_doc(df.loc[440].content)

"""*** preproces example for one doc***"""

normalized_doc= Normalizing_doc(df.loc[440].content)
normalized_doc

spellchecked_doc= SpellChecking_doc(normalized_doc)
spellchecked_doc

tokens= Tokenizing_doc(spellchecked_doc)
tokens

cleaned_token=remove_uselesstokens(tokens)
cleaned_token

stemmed_words=Stemming_doc(cleaned_token)
stemmed_words

cleaned_words=remove_stopwords(stemmed_words)
cleaned_words

print(len(tokens))
print( len(cleaned_token))
print( len(stemmed_words))
print( len(cleaned_words))

sorted([[x,cleaned_words.count(x)] for x in set(cleaned_words)],reverse=True, key = lambda x: x[1])

df

"""**Spiliting Data**"""

x = df.index
y = df['label']
x

x_train, x_temp, y_train, y_temp = train_test_split(x,y, test_size=.15,stratify=y)

x_validation,x_test ,y_validation,y_test=train_test_split(x_temp,y_temp, test_size=.05,stratify=y_temp)

cc,x_test ,ccc,y_test=train_test_split(x_temp,y_temp, test_size=.5,stratify=y_temp)

print(len(x_train))
print(len(x_validation))
print(len(x_test))

mydata_id= list(x_train)+ list( x_validation )+list( x_test )

len( mydata_id)

"""***preprocessing for a group of docs***"""

import sys
print(sys.getrecursionlimit())

sys.setrecursionlimit(5000)

mydata=[]
counter=1
for d in mydata_id:
    print('processing doc :'+ str(counter))
    normalized_doc= Normalizing_doc(df.loc[d].content)
    spellchecked_doc= SpellChecking_doc(normalized_doc)
    tokens= Tokenizing_doc(spellchecked_doc)
    cleaned_token=remove_uselesstokens(tokens)
    stemmed_words=Stemming_doc(cleaned_token)
    mydata.append( remove_stopwords(stemmed_words))
    counter=counter+1

mydata[50]

"""# 5-Feature Engineering"""

for i in range(len( mydata )):
    sd= sorted([[x,mydata[i].count(x)] for x in set(mydata[i])],reverse=True, key = lambda x: x[1])
    print('doc :'+ str(i) )
    print(sd[0:20])

"""Wordnet connection"""
'''
Base = automap_base()
engine = create_engine('sqlite:////content/drive/MyDrive/farsnet/farsnet2.5.db3')
Base.prepare(engine, reflect=True)

word_table = Base.classes.word
session = Session(engine)

session.query(word_table)

"""**test connection**"""

qur="واقعا".replace('ی','ي')
myword =session.query(word_table).filter_by(defaultValue=qur).first()
if myword is None:
  print('does not exist')
else:
  print( myword.defaultValue)

"""**dictionary forming**"""

mydic=[]
for r in session.query(word_table).all():
  r.defaultValue=r.defaultValue.replace('ی','ي')
  mydic.append(r.defaultValue)

len(mydic)

rr="باستانی".replace('ی','ي')
rr in mydic

"""**Dividing words into match & mismatch**"""

mismatches=[]
matches=[]
# for d in mydata:
#     mismatches.append( [w for w in d if w.replace('ی','ي') not in mydic])
#     matches.append( [w for w in d if w.replace('ی','ي')  in mydic])

mismatches.append( [w for w in cleaned_words if w.replace('ی','ي') not in mydic])
matches.append( [w for w in cleaned_words if w.replace('ی','ي')  in mydic])
print('mismatch :'+str( len(mismatches[0])) )
print('match :'+str(len(matches[0])) )

# for i in range(len(stemmed_words)):
#     print('doc'+str(i)+':--------------------------')
#     print('total words :'+ str(len(stemmed_words[i])) )
#     print('mismatch :'+str( len(mismatches[i])) )
#     print('match :'+str(len(matches[i])) )
# mismatches

mismatches[0]

matches[0]
'''
"""**TF-IDF**"""

corpus_train=[]
for d in mydata[0 : len(x_train)]:
  corpus_train.append( " ".join(d))


corpus_validation=[]
for d in mydata[len(x_train) : len(x_train)+len(x_validation)]:
  corpus_validation.append( " ".join(d))


corpus_test=[]
for d in mydata[len(x_train)+len(x_validation) : len(mydata_id)]:
  corpus_test.append( " ".join(d))

corpus_train[10]

"""feature extraction"""
mxf=300
xtrain_tfidf=[]
xvalid_tfidf=[]
while mxf<500 :
    vectorizer = TfidfVectorizer(max_features=mxf)
    xtrain_tfidf.append( vectorizer.fit_transform(corpus_train))
    xvalid_tfidf.append( vectorizer.fit_transform(corpus_validation))
    mxf=mxf+5
print(vectorizer.get_feature_names())

"""# 6-Model Building"""

def train_model(classifier, feature_vector_train, label, feature_vector_valid, is_neural_net=False):
    # fit the training dataset on the classifier
    classifier.fit(feature_vector_train, label)

    # predict the labels on validation dataset
    predictions = classifier.predict(feature_vector_valid)

    if is_neural_net:
        predictions = predictions.argmax(axis=-1)
    return predictions

naive_pred=[]
log_pred=[]
for_pred=[]
xg_pred=[]
for i in range(len( xtrain_tfidf)):
  naive_pred.append(train_model(MultinomialNB(), xtrain_tfidf[i], y_train.to_list(), xvalid_tfidf[i]))
  log_pred.append( train_model(LogisticRegression(), xtrain_tfidf[i], y_train.to_list(), xvalid_tfidf[i]))
  for_pred.append( train_model(RandomForestClassifier(), xtrain_tfidf[i], y_train.to_list(), xvalid_tfidf[i]))
  xg_pred.append( train_model( xgboost.XGBClassifier(),xtrain_tfidf[i].tocsc(), y_train.to_list(), xvalid_tfidf[i].tocsc()))
"""#7-Performance Measurement"""

def metricModel(y_real,y_pred):
  print(metrics.confusion_matrix(y_real, y_pred))
  print(metrics.classification_report(y_real, y_pred))

# Naive bayse
a=[]
for i in range(len( naive_pred)):
  a.append( metrics.accuracy_score(naive_pred[i], y_validation.to_list()))
print("best model of naive bayse :\n ")  
metricModel(y_validation.to_list(),naive_pred[a.index(max(a))])

# Logistic Regression
a=[]
for i in range(len( log_pred)):
  a.append( metrics.accuracy_score(log_pred[i], y_validation.to_list()))
print("best model of LogisticRegression :\n ")
metricModel(y_validation.to_list(),log_pred[a.index(max(a))])

# Random Forest
a=[]
for i in range(len( for_pred)):
  a.append( metrics.accuracy_score(for_pred[i], y_validation.to_list()))
print("best model of RandomForest :\n ")
metricModel(y_validation.to_list(),for_pred[a.index(max(a))])

# Xgboost
a=[]
for i in range(len( xg_pred)):
  a.append( metrics.accuracy_score(xg_pred[i], y_validation.to_list()))
print("best model of xgboost :\n ")
metricModel(y_validation.to_list(),xg_pred[a.index(max(a))])

